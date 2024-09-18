import pandas as pd
from io import BytesIO
import markdown2
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
from datetime import datetime

def add_html_to_docx(soup, doc: Document):
    """
    Add HTML content to Word document using docx.
    
    Args:
        soup (BeautifulSoup): Parsed HTML content
        doc (Document): Word document object
    """
    # Iterate through HTML elements and add to document
    for element in soup:
        if element.name == 'h1':
            doc.add_heading(element.text, level=1)
        elif element.name == 'h2':
            doc.add_heading(element.text, level=2)
        elif element.name == 'p':
            doc.add_paragraph(element.text)
        elif element.name == 'ul':
            for li in element.find_all('li'):
                doc.add_paragraph(f'• {li.text}', style='List Bullet')

def add_dataframe_to_docx(doc: Document, df: pd.DataFrame):
    """
    Add pandas DataFrame as table to Word document.
    
    Args:
        doc (Document): Word document object
        df (pd.DataFrame): DataFrame to be added
    """

    rows, cols = df.shape
    t = doc.add_table(rows + 1, cols)
    t.style = 'Table Grid'

    headers = df.columns.to_numpy()
    data = df.values

    # Aggiungi le intestazioni
    for j, header in enumerate(headers):
        t.cell(0, j).text = str(header)

    # Aggiungi il resto dei dati
    for i in range(rows):
        row = data[i]
        for j, value in enumerate(row):
            t.cell(i + 1, j).text = str(value)

# Function to generate Word report
def generate_word_report(dfs = {}, analyses = {}, figures = {}):
    """
    Generate Word report with analyses, figures, and DataFrames.
    
    Args:
        dfs (dict): Dictionary of DataFrames
        analyses (dict): Dictionary of markdown-formatted analyses
        figures (dict): Dictionary of plotly figures
    
    Returns:
        BytesIO: Word document as bytes object
    """
    doc = Document()
    doc.add_heading('Orizon Security Dashboard Report', 0)

    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_page_break()

    # Table of Contents
    doc.add_heading('Table of Contents', level=1)
    for section in analyses.keys():
        doc.add_paragraph(section.capitalize(), style='List Bullet')
    doc.add_page_break()

    for section, content in analyses.items():
        doc.add_heading(section.capitalize(), level=1)
        # Converte markdown in HTML
        html = markdown2.markdown(content)
        soup = BeautifulSoup(html, 'html.parser')
        add_html_to_docx(soup, doc)

        if section in figures:
            for fig in figures[section]:
                img_buffer = BytesIO()
                try:
                    fig.write_image(img_buffer, format="png")
                    doc.add_picture(img_buffer, width=Inches(5))
                except:
                    pass

    for key, df in dfs.items():
        doc.add_heading(key.capitalize(), level=1)
        add_dataframe_to_docx(doc, df)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer