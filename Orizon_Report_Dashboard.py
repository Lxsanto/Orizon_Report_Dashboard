import os
import io
import zipfile
import time
import subprocess
from collections import Counter
from datetime import datetime, timedelta
import pandas as pd
import streamlit as st
import plotly.express as px
import plotly.io as pio
import networkx as nx
from io import BytesIO
from wordcloud import WordCloud
import torch
import cProfile
import pstats
from matplotlib.colors import LinearSegmentedColormap

# Docx per la gestione di documenti Word
from docx import Document
from docx.shared import Inches

# Streamlit Authenticator per la gestione dell'autenticazione
from streamlit_authenticator import Authenticate
import yaml
from yaml.loader import SafeLoader

# my functions
from restart_utils import clear_pycache, restart_script
from GPU_utils import print_gpu_utilization, print_summary
from graph_utils import *  # all Michele utils
from prompts_utils import *

# Tentativo di importazione condizionale con gestione degli errori
try:
    from transformers import AutoTokenizer, pipeline
except ImportError:
    print("Errore durante l'importazione di 'AutoTokenizer'. Pulizia della cache e riavvio dello script...")
    clear_pycache()
    restart_script()

logo = Image.open("logo1.png")

# Definizione dei colori del branding kit
kelly_green = "#4AC300"
mariana_blue = "#002430"
burnt_red = "#E5625E"
dodger_blue = "#2191FB"
dawn_mist = "#DBE2E9"
simple_white = "#FFFFFF"
sunglow = "#FFC857"

# Configurazione di Streamlit
st.set_page_config(page_title="Orizon Security", layout="wide", initial_sidebar_state="expanded")
with st.sidebar:
    st.image(logo)
st.config.set_option("theme.base", "light")
st.config.set_option("theme.primaryColor", kelly_green)
st.config.set_option("theme.backgroundColor", simple_white)
st.config.set_option("theme.secondaryBackgroundColor", dawn_mist)
st.config.set_option("theme.textColor", mariana_blue)
st.config.set_option("theme.font", 'sans serif')

# Configurazione dell'autenticazione Streamlit
with open('password.yaml') as file:
    config = yaml.load(file, Loader=SafeLoader)

authenticator = Authenticate(
    config['credentials'],
    config['cookie']['name'],
    config['cookie']['key'],
    config['cookie']['expiry_days'],
    config['preauthorized']
)

# Configurazione di Plotly
template = pio.templates['ggplot2']
pio.templates.default = 'ggplot2'

template.layout.font.family = "Gill Sans, sans-serif"
template.layout.font.size = 12  # Ridotto da 300 a una dimensione più ragionevole
template.layout.font.color = mariana_blue
template.layout.title.font.size = 20
template.layout.xaxis.title.font.size = 16
template.layout.yaxis.title.font.size = 16
template.layout.paper_bgcolor = simple_white
template.layout.plot_bgcolor = dawn_mist  # Cambiato da rosso a un colore più neutro
template.layout.xaxis.gridcolor = simple_white
template.layout.xaxis.linecolor = mariana_blue
template.layout.xaxis.tickcolor = mariana_blue
template.layout.yaxis.gridcolor = simple_white
template.layout.yaxis.linecolor = mariana_blue
template.layout.yaxis.tickcolor = mariana_blue

# Definizione della palette di colori personalizzata
colors = [kelly_green, dodger_blue, burnt_red, mariana_blue]
template.layout.colorway = colors

pio.templates["Orizon_template"] = template
pio.templates.default = "Orizon_template"
_width = 800 
_height = 600

# Configurazione dell'ambiente CUDA
are_you_on_CUDA = False
run_LLM = False
if are_you_on_CUDA:
    os.environ['PYTORCH_CUDA_ALLOC_CONF'] = 'expandable_segments:True'

# Selezione del modello LLM
#model_id = "Qwen/Qwen2-1.5B-Instruct"
model_id = "microsoft/Phi-3.5-mini-instruct"
auth_token = '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'

# Pulizia della cache di Streamlit
clear = False
if clear:
    st.cache_resource.clear()

### Streamlit CSS dashboard setups ###
st.markdown("""
                <style>
                @import url('https://fonts.googleapis.com/css2?family=Gill+Sans&display=swap');

                body {
                    font-family: 'Gill Sans', sans-serif;
                    background-color: #FFFFFF;
                    color: #002430;
                }

                h1, h2, h3 {
                    color: #002430;
                    font-weight: 600;
                }

                .stAlert {
                    background-color: #DBE2E9;
                    color: #002430;
                    border: none;
                    border-radius: 8px;
                    box-shadow: 0 2px 4px rgba(0, 0, 0, 0.05);
                }

                .stButton > button {
                    background-color: #2191FB;  /* Dodger Blue */
                    color: #FFFFFF;  /* Simple White */
                    border: none;
                    border-radius: 4px;
                    padding: 0.5rem 1rem;
                    font-weight: 600;
                }

                .stButton > button:hover {
                    background-color: #4AC300;  /* Kelly Green */
                    color: #FFFFFF;  /* Simple White */
                }

                .stProgress .st-bo {
                    background-color: #4AC300;
                }

                div[data-testid="stMetricValue"] {
                    font-size: 2rem;
                    font-weight: 600;
                    color: #002430;
                }

                div[data-testid="stMetricLabel"] {
                    font-size: 0.9rem;
                    font-weight: 400;
                    color: #002430;
                }

                .plot-container {
                    border-radius: 8px;
                    box-shadow: 0 2px 4px rgba(0, 0, 0, 0.05);
                    background-color: #FFFFFF;
                    padding: 1rem;
                    margin-bottom: 1.5rem;
                }

                .stSelectbox, .stMultiSelect {
                    background-color: #FFFFFF;
                    color: #002430;
                    border-radius: 6px;
                    box-shadow: 0 1px 2px rgba(0, 0, 0, 0.05);
                }

                .sidebar .sidebar-content {
                    background-color: #FFFFFF;
                }

                .tooltip {
                    position: relative;
                    display: inline-block;
                    border-bottom: 1px dotted #002430;
                }

                .tooltip .tooltiptext {
                    visibility: hidden;
                    width: 200px;
                    background-color: #002430;
                    color: #FFFFFF;
                    text-align: center;
                    border-radius: 6px;
                    padding: 5px 0;
                    position: absolute;
                    z-index: 1;
                    bottom: 125%;
                    left: 50%;
                    margin-left: -100px;
                    opacity: 0;
                    transition: opacity 0.3s;
                }

                .tooltip:hover .tooltiptext {
                    visibility: visible;
                    opacity: 1;
                }

                .kpi-card {
                    background-color: #FFFFFF;
                    border-radius: 8px;
                    box-shadow: 0 2px 4px rgba(0, 0, 0, 0.1);
                    padding: 1rem;
                    text-align: center;
                }

                .kpi-value {
                    font-size: 2rem;
                    font-weight: 600;
                    color: #4AC300;
                }

                .kpi-label {
                    font-size: 0.9rem;
                    color: #002430;
                }

                .section-nav {
                    position: sticky;
                    top: 0;
                    background-color: #FFFFFF;
                    z-index: 1000;
                    padding: 1rem 0;
                    margin-bottom: 1rem;
                    border-bottom: 1px solid #002430;
                }

                .section-nav a {
                    color: #002430;
                    text-decoration: none;
                    margin-right: 1rem;
                    padding: 0.5rem 1rem;
                    border-radius: 4px;
                    transition: background-color 0.2s ease;
                }

                .section-nav a:hover {
                    background-color: #2191FB;
                    color: #FFFFFF;
                }
                </style>
    """, unsafe_allow_html=True)

### Utility functions ###
@st.cache_resource
def load_LLM(model_id = model_id, auth_token = auth_token):

    if not torch.cuda.is_available():
        print("CUDA GPU not available!")
    else:
        print("CUDA is available. Backend and pinned memory configurations are applied.")
        print_gpu_utilization()

    
    try:
        model_kwargs = {
        'attn_implementation': "flash_attention_2"
        #"torch_dtype": torch.bfloat16
        }
        _tokenizer = AutoTokenizer.from_pretrained(model_id, 
                                                   use_fast= True)
        chat_pipeline = pipeline("text-generation", 
                                 torch_dtype='auto',
                             model=model_id,
                             token=auth_token,
                             tokenizer=_tokenizer,
                             #device='cuda:0',
                             device_map="auto",
                             #model_kwargs=model_kwargs
                             )
        
        print(f'Pipeline loaded on {chat_pipeline.device}')

        return chat_pipeline
    
    except Exception as e:
        st.error(f"Error loading the model: {str(e)}")


@st.cache_data
def calculate_risk_score(vulnerabilities, severity_column):
    # Severity weights
    severity_weights = {'critical': 80, 'high': 60, 'medium': 40, 'low': 10, 'info': 1}
    
    # Count vulnerabilities by severity
    severity_counts = {severity: 0 for severity in severity_weights}
    for v in vulnerabilities[severity_column]:
        severity = str(v).lower()
        if severity in severity_counts:
            severity_counts[severity] += 1
    
    # Calculate weighted score
    weighted_score = sum(severity_weights[severity] * count for severity, count in severity_counts.items())
    
    # Calculate risk score
    total_vulnerabilities = sum(severity_counts.values())
    if total_vulnerabilities == 0:
        return 0
    
    # Base the score primarily on the weighted score, but consider the number of vulnerabilities
    risk_score = (weighted_score / total_vulnerabilities) * 2
    
    # Ensure medium vulnerabilities have a significant impact
    if severity_counts['medium'] > 0:
        risk_score = max(risk_score, 40 + (severity_counts['medium'] - 1) * 5)
    
    # Cap the score at 100
    risk_score = min(risk_score, 100)
    
    return int(risk_score)

@st.cache_data
def create_severity_impact_bubble(vulnerabilities, severity_column, cvss_column, host_column):
    if all(col in vulnerabilities.columns for col in [severity_column, cvss_column, host_column]):
        vulnerability_counts = vulnerabilities.groupby([severity_column, host_column]).size().reset_index(name='count')
        avg_cvss = vulnerabilities.groupby([severity_column, host_column])[cvss_column].mean().reset_index(name='avg_cvss')
        bubble_data = pd.merge(vulnerability_counts, avg_cvss, on=[severity_column, host_column])
        
        fig = px.scatter(bubble_data, 
                         width=_width,
                         height=_height,
                         x='count', 
                         y='avg_cvss', 
                         size='count', 
                         color=severity_column,
                         hover_name=host_column,
                         labels={'count': 'Number of Vulnerabilities', 'avg_cvss': 'Average CVSS Score'},
                         title="Severity, Impact, and Prevalence Correlation")
        return fig
    else:
        return None

# Function to generate Word report
def generate_word_report(vulnerabilities, analyses, figures):
    doc = Document()
    doc.add_heading('Orizon Security Dashboard Report', 0)

    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_page_break()

    # Table of Contents
    doc.add_heading('Table of Contents', level=1)
    for section in analyses.keys():
        doc.add_paragraph(section.capitalize(), style='List Bullet')
    doc.add_page_break()

    for section, content in analyses.items():
        doc.add_heading(section.capitalize(), level=1)
        doc.add_paragraph(content)
        
        if section in figures:
            img_buffer = BytesIO()
            figures[section].write_image(img_buffer, format="png", width=800, height=400, scale=2)
            doc.add_picture(img_buffer, width=Inches(7.5))

    # Add summary tables
    doc.add_heading('Vulnerability Summary', level=1)
    
    # Severity Distribution Table
    severity_counts = vulnerabilities['severity'].value_counts()
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Severity'
    hdr_cells[1].text = 'Count'
    hdr_cells[2].text = 'Percentage'
    for severity, count in severity_counts.items():
        percentage = (count / len(vulnerabilities)) * 100
        row_cells = table.add_row().cells
        row_cells[0].text = severity
        row_cells[1].text = str(count)
        row_cells[2].text = f"{percentage:.2f}%"

    doc.add_paragraph()

    # Top 10 Vulnerabilities Table
    doc.add_heading('Top 10 Vulnerabilities', level=2)
    top_10 = vulnerabilities.sort_values('severity', ascending=False).head(10)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Host'
    hdr_cells[1].text = 'Severity'
    hdr_cells[2].text = 'Vulnerability'
    hdr_cells[3].text = 'Description'
    for _, row in top_10.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = row['host']
        row_cells[1].text = row['severity']
        row_cells[2].text = row['template_name']
        row_cells[3].text = row['description'][:50] + '...'

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

@st.cache_data
def process_and_filter_vulnerabilities(uploaded_file):
    vulnerabilities = load_data(uploaded_file)

    if vulnerabilities is not None and not vulnerabilities.empty:
        st.sidebar.success("JSON file loaded successfully!")
        
        if 'created_at' in vulnerabilities.columns:
            vulnerabilities['created_at'] = pd.to_datetime(vulnerabilities['created_at'], errors='coerce')
            vulnerabilities = vulnerabilities.dropna(subset=['created_at'])
            return vulnerabilities
        else:
            st.error("The 'created_at' column is missing from the data. Please check your JSON file.")
            return None
    else:
        st.error("Failed to load data or the file is empty.")
        return None

@st.cache_data
def create_risk_score_gauge(risk_score):

    # Determinazione del colore del gauge basato sul risk_score
    if 20 < risk_score < 60:
        gauge_color = sunglow
    elif risk_score >= 60:
        gauge_color = burnt_red
    else:
        gauge_color = kelly_green
    
    # Creazione della figura
    fig = go.Figure(
        go.Indicator(
            mode="gauge+number",
            value=risk_score,
            domain={'x': [0, 1], 'y': [0, 1]},
            title={'text': "Risk Score", 'font': {'size': 20}},
            gauge={
                'bar': {'color': gauge_color},
                'axis': {'range': [0, 100], 'tickwidth': 1, 'tickcolor': mariana_blue}
            }
        ),
        layout=go.Layout(
            width=_width,
            height=_height,
            font={'color': mariana_blue}
        )
    )
    
    return fig

@st.cache_data
def pie(severity_counts):

    fig_severity = go.Figure(data=[go.Pie(
        labels=severity_counts.index,
        values=severity_counts.values,
        textinfo='percent+label',
        textposition='inside',
        hole=0.3,
        pull=[0.1] * len(severity_counts),  # This creates the exploded effect
        marker=dict(colors=severity_counts.index),  # Use the same colors as before
    )])

    fig_severity.update_layout(
        title_text="Vulnerability Severity Distribution",
        title_x=0.5,  # Center the title
        width=_width,
        height=_height,
        scene=dict(
            xaxis_title='',
            yaxis_title='',
            zaxis_title='',
            aspectmode='manual',
            aspectratio=dict(x=1, y=1, z=0.5)  # This gives a 3D effect
        ),
        showlegend=True,
        legend=dict(orientation="h", yanchor="bottom", y=-0.1, xanchor="center", x=0.5)
    )

    fig_severity.update_traces(
        textfont_size=12,
        marker=dict(line=dict(color='#000000', width=2))  # Add a black outline to each slice
    )

    return fig_severity

def Geolocation_of_servers(file_contents, api_key):
    # Load and preprocess data
    df = load_data_geo(file_contents)
    
    severity_weights = {
        'unknown': 1, 'info': 2, 'low': 4, 'medium': 6, 'high': 8, 'critical': 10
    }
    
    df['severity_weight'] = df['severity'].map(severity_weights)
    danger_score_per_server = df.groupby('host')['severity_weight'].sum().reset_index()
    host_names = danger_score_per_server['host']

    st.write('The geolocation process duration depends on DNS settings... Please wait')
    progress_bar = st.progress(0)
    status_text = st.empty()
    summary = st.empty()

    geo_results = []
    ips = []
    for id, host in enumerate(host_names):
        progress = (id + 1) / len(host_names)
        progress_bar.progress(progress)
        status_text.text(f'Numbers of scanned hosts: {id + 1}/{len(host_names)}')
        ip = resolve_hostname(host)
        ips.append(ip)
        d = geolocate_ip(ip, api_key)
        geo_results.append(d)


    danger_score_per_server['ip'] = ips
    
    geolocation_data = pd.DataFrame(geo_results, columns=['latitude', 'longitude', 'country', 'city'])
    
    # Combine data and aggregate risk scores
    risk_by_ip = pd.concat([danger_score_per_server, geolocation_data], axis=1)
    risk_by_ip = risk_by_ip.groupby(['ip', 'country', 'city', 'latitude', 'longitude'])['severity_weight'].sum().reset_index()
    
    # Normalize risk scores
    max_score = risk_by_ip['severity_weight'].max()
    risk_by_ip['normalized_risk_score'] = (risk_by_ip['severity_weight'] / max_score) * 100
    risk_by_ip.loc[risk_by_ip['severity_weight'] == max_score, 'normalized_risk_score'] = 100

    geo_map = create_plotly_map(risk_by_ip)
    geo_map_1 = create_country_bubble_plot(risk_by_ip)

    # Group hosts by IP
    hosts_by_ip = danger_score_per_server.groupby('ip')['host'].agg(list).reset_index()
    hosts_by_ip.columns = ['ip', 'associated_hosts']

    # Merge the hosts information with the risk_by_ip dataframe
    risk_by_ip = risk_by_ip.merge(hosts_by_ip, on='ip', how='left')
    
    return geo_map, geo_map_1, risk_by_ip

def main():
    st.sidebar.title("Orizon Security Dashboard")

    if st.sidebar.button("Restart App"):
        subprocess.run(['python', 'run_streamlit_port8501.py'])
        print('Dashboard is now restarted!')
    
    language = st.selectbox('select language here',
                            ('English', 'Italian', 'Spanish'))
    st.write('You selected:', language)
    
    if language == 'Italian':
        language = 'it'
    if language == 'English':
        language = 'en'
    if language == 'Spanish':
        language = 'es'
    
    uploaded_file = st.sidebar.file_uploader("Upload Vulnerability JSON", type="json", key="vuln_upload")
    
    if uploaded_file:

        filtered_vulnerabilities = process_and_filter_vulnerabilities(uploaded_file)

        # Main content
        st.title("Orizon Security Dashboard")
        st.markdown("Welcome to our private Security Dashboard, here you can see the analysis of the JSON file.")

        pipe = None
        # Load model
        if run_LLM:
            pipe = load_LLM()

        # Automatic column detection
        severity_column = 'severity' if 'severity' in filtered_vulnerabilities.columns else None
        description_column = 'description' if 'description' in filtered_vulnerabilities.columns else None
        created_at_column = 'created_at' if 'created_at' in filtered_vulnerabilities.columns else None
        host_column = 'host' if 'host' in filtered_vulnerabilities.columns else None

        # Navigation
        st.markdown("""
        <div class="section-nav">
            <a href="#security-posture-overview">Overview</a>
            <a href="#vulnerability-severity-distribution">Severity</a>
            <a href="#top-10-critical-vulnerabilities">Top Vulnerabilities</a>
            <a href="#network-topology-analysis">Network Analysis</a>
            <a href="#additional-cybersecurity-insights">Additional Insights</a>
        </div>
        """, unsafe_allow_html=True)

        # Executive Summary
        st.header("Executive Summary", anchor="executive-summary")
        total_vulns = len(filtered_vulnerabilities)
        risk_score = calculate_risk_score(filtered_vulnerabilities, severity_column)
        critical_vulns = len(filtered_vulnerabilities[filtered_vulnerabilities[severity_column].str.lower() == 'critical'])
        high_vulns = len(filtered_vulnerabilities[filtered_vulnerabilities[severity_column].str.lower() == 'high'])
        medium_vulns = len(filtered_vulnerabilities[filtered_vulnerabilities[severity_column].str.lower() == 'medium'])
        low_vulns = len(filtered_vulnerabilities[filtered_vulnerabilities[severity_column].str.lower() == 'low'])
        
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.markdown(f"""
            <div class="kpi-card">
                <div class="kpi-value">{total_vulns}</div>
                <div class="kpi-label">Total Vulnerabilities</div>
            </div>
            """, unsafe_allow_html=True)
        with col2:
            st.markdown(f"""
            <div class="kpi-card">
                <div class="kpi-value">{risk_score}/100</div>
                <div class="kpi-label">Risk Score</div>
            </div>
            """, unsafe_allow_html=True)
        with col3:
            st.markdown(f"""
            <div class="kpi-card">
                <div class="kpi-value">{critical_vulns}</div>
                <div class="kpi-label">Critical Vulnerabilities</div>
            </div>
            """, unsafe_allow_html=True)
        with col4:
            st.markdown(f"""
            <div class="kpi-card">
                <div class="kpi-value">{high_vulns}</div>
                <div class="kpi-label">High Vulnerabilities</div>
            </div>
            """, unsafe_allow_html=True)

        # Security Posture Overview
        st.header("Security Posture Overview", anchor="security-posture-overview")
        col1, col2 = st.columns([3, 2])
        with col1:

            fig_risk_score = create_risk_score_gauge(risk_score)
            st.plotly_chart(fig_risk_score, use_container_width=True, config={'displayModeBar': False})
        
        with col2:
            st.subheader("Orizon Engine Analysis")
            overview_analysis = ''
            with st.spinner("Generating overview analysis..."):
                if run_LLM:
                    overview_analysis = analyze_overview(total_vulns, risk_score, critical_vulns, high_vulns, medium_vulns, low_vulns, _pipe = pipe, language=language)
            st.markdown(overview_analysis)

        # Severity Distribution
        st.header("Vulnerability Severity Distribution", anchor="vulnerability-severity-distribution")
        col1, col2 = st.columns([2, 1])
        with col1:
            severity_counts = filtered_vulnerabilities[severity_column].value_counts()
            fig_severity = pie(severity_counts)

            st.plotly_chart(fig_severity, use_container_width=True, config={'displayModeBar': False})

        with col2:
            st.subheader("Orizon Engine Analysis")
            severity_analysis = ''
            with st.spinner("Generating severity analysis..."):
                if run_LLM:
                    severity_analysis = analyze_severity_distribution(severity_counts, _pipe= pipe, language=language)
            st.markdown(severity_analysis)


        # Geolocation of servers
        st.header("Geolocation of company servers", anchor="Geolocation of company servers")
        
        col1, col2 = st.columns([2, 1])
        with col1:

            file_contents = uploaded_file.read()

            with cProfile.Profile() as pr:
                geo_map, geo_map_1, risk_by_ip = Geolocation_of_servers(file_contents, api_key='f2cfc8c5c8c358')
            with open("profiling_results_geo.txt", "w") as f:
                stats = pstats.Stats(pr, stream=f)
                stats.sort_stats('cumulative')
                stats.print_stats()
            
            # Create and display the Plotly maps
            st.plotly_chart(geo_map, use_container_width=True, config={'displayModeBar': False})
            st.plotly_chart(geo_map_1, use_container_width=True, config={'displayModeBar': False})

            # Display the data in the table
            st.subheader("Risk Scores by IP")

            # Update the selected columns to include the new 'associated_hosts' column
            selected_columns = ['ip', 'associated_hosts', 'country', 'city', 'severity_weight', 'normalized_risk_score']

            # Pagination
            items_per_page = st.slider("Items per page", min_value=10, max_value=100, value=20, step=10)
            total_pages = len(risk_by_ip) // items_per_page + (1 if len(risk_by_ip) % items_per_page > 0 else 0)
            current_page = st.number_input("Page", min_value=1, max_value=total_pages, value=1)

            start_idx = (current_page - 1) * items_per_page
            end_idx = start_idx + items_per_page

            # Display the selected page of the table
            st.dataframe(risk_by_ip[selected_columns].iloc[start_idx:end_idx], height=400, use_container_width=True)

            # Show the pagination information
            st.write(f"Showing {start_idx+1} to {min(end_idx, len(risk_by_ip))} of {len(risk_by_ip)} entries")
        
        #with col2:
            #with st.spinner("Generating analysis..."):
                #if run_LLM:
                    #geo_analysis = analyze_geolocation(ip = risk_by_ip['ip'], _pipe = pipe, language=language)
                    #st.markdown(overview_analysis)

        # Top 10 Vulnerabilities
        st.header("Top 10 Critical Vulnerabilities", anchor="top-10-critical-vulnerabilities")
    
        def severity_to_num(severity):
            severity_order = {'Critical': 6, 'High': 5, 'Medium': 4, 'Low': 3, 'Info': 2, 'Unknown': 1}
            return severity_order.get(severity.capitalize(), 0)

        # Apply the custom sorting
        filtered_vulnerabilities['severity_num'] = filtered_vulnerabilities[severity_column].apply(severity_to_num)
        top_10 = filtered_vulnerabilities.sort_values(['severity_num', severity_column], ascending=[False, False]).head(10)
    
    
        # Apply the custom sorting
        filtered_vulnerabilities['severity_num'] = filtered_vulnerabilities[severity_column].apply(severity_to_num)
        sorted_vulnerabilities = filtered_vulnerabilities.sort_values(['severity_num', severity_column], ascending=[False, False])

        # Select the columns to display
        selected_columns = [host_column, severity_column, 'template_name', description_column]

        # Pagination
        severities_page = st.slider("Severities per page", min_value=10, max_value=100, value=20, step=10)
        total_pages = len(sorted_vulnerabilities) // severities_page + (1 if len(sorted_vulnerabilities) % severities_page > 0 else 0)
        current_page = st.number_input("Page", min_value=1, max_value=total_pages, value=1)

        start_idx = (current_page - 1) * severities_page
        end_idx = start_idx + severities_page

        # Display the selected page of the table
        st.dataframe(sorted_vulnerabilities[selected_columns].iloc[start_idx:end_idx], height=400, use_container_width=True)

        # Show the pagination information
        st.write(f"Showing {start_idx+1} to {min(end_idx, len(sorted_vulnerabilities))} of {len(sorted_vulnerabilities)} entries")
        st.subheader("Orizon Engine Analysis")
        common_types = top_10['template_name'].value_counts()
        most_common_type = common_types.index[0]
        hosts_affected = top_10[host_column].nunique()
        most_affected_host = top_10[host_column].value_counts().index[0]
        top_vuln_analysis = ''
        with st.spinner("Analyzing top vulnerabilities..."):
            top_vuln_analysis = ''
            if run_LLM:
                top_vuln_analysis = analyze_top_vulnerabilities(most_common_type, common_types, hosts_affected, most_affected_host, _pipe = pipe, language=language)
        st.markdown(top_vuln_analysis)

        # Network Topology View
        st.header("Network Topology Analysis", anchor="network-topology-analysis")
        G = nx.Graph()
        for _, row in filtered_vulnerabilities.iterrows():
            G.add_edge(row[host_column], row['template_name'])
        pos = nx.spring_layout(G)
        edge_x, edge_y = [], []
        for edge in G.edges():
            x0, y0 = pos[edge[0]]
            x1, y1 = pos[edge[1]]
            edge_x.extend([x0, x1, None])
            edge_y.extend([y0, y1, None])
        edge_trace = go.Scatter(x=edge_x, y=edge_y, line=dict(width=0.5), hoverinfo='none', mode='lines')
        node_x = [pos[node][0] for node in G.nodes()]
        node_y = [pos[node][1] for node in G.nodes()]
        node_trace = go.Scatter(x=node_x, y=node_y, mode='markers', hoverinfo='text',
                                marker=dict(showscale=True, size=10, 
                                            color=[], colorbar=dict(thickness=15, title='Node Connections'),
                                            line_width=2))
        node_adjacencies = []
        node_text = []
        for node, adjacencies in enumerate(G.adjacency()):
            node_adjacencies.append(len(adjacencies[1]))
            node_text.append(f'{adjacencies[0]} - # of connections: {len(adjacencies[1])}')
        node_trace.marker.color = node_adjacencies
        node_trace.text = node_text
        fig_network = go.Figure(data=[edge_trace, node_trace],
                        layout=go.Layout(showlegend=False, hovermode='closest',
                                         title="Network Topology Visualization", width=_width, height=_height))
        st.plotly_chart(fig_network, use_container_width=True, config={'displayModeBar': False})
        st.subheader("Orizon Engine Analysis")
        centrality = nx.degree_centrality(G)
        top_central = sorted(centrality, key=centrality.get, reverse=True)[:5]
        density = nx.density(G)
        communities = list(nx.community.greedy_modularity_communities(G))
        #with st.spinner("Analyzing network topology..."):
        #    if run_LLM:
        #        network_analysis = generate_network_analysis(top_central, density, communities, _pipe=pipe, language=language)
        #        st.markdown(network_analysis)

        # Additional Cybersecurity Insights
        st.header("Additional Cybersecurity Insights", anchor="additional-cybersecurity-insights")
        
        # CVSS Score Distribution (if available)
        if 'cvss_score' in filtered_vulnerabilities.columns:
            st.subheader("CVSS Score Distribution")
            fig_cvss = px.histogram(
                filtered_vulnerabilities,
                width=_width,
                height=_height, 
                x='cvss_score', 
                nbins=20, 
                title="Distribution of CVSS Scores",
                labels={'cvss_score': 'CVSS Score', 'count': 'Number of Vulnerabilities'}
            )
            fig_cvss.update_layout(bargap=0.1)
            st.plotly_chart(fig_cvss, use_container_width=True, config={'displayModeBar': False})
            
            avg_cvss = filtered_vulnerabilities['cvss_score'].mean()
            high_cvss = filtered_vulnerabilities[filtered_vulnerabilities['cvss_score'] > 7]
            cvss_analysis = ''
            with st.spinner("Analyzing CVSS distribution..."):
                cvss_analysis = ''
                if run_LLM:
                    cvss_analysis = analyze_cvss_distribution(avg_cvss, len(high_cvss), total_vulns, _pipe = pipe, language=language)
            st.markdown(cvss_analysis)

        if created_at_column:
            # Michele
            st.subheader("Screenshots")
            st.write("We are taking screenshots...")

            scelta = 'No'
            #if st.button('Click here to interrupt the process'):
            #   scelta = 'Yes'

            if scelta == 'No':

                screenshots = []
                errors = []
                progress_bar = st.progress(0)
                status_text = st.empty()
                summary = st.empty()

                df = load_data_screen(file_contents)

                # Filter the dataframe
                filtered_df = df[~df['severity'].isin(['info'])]

                # Get unique hosts
                unique_hosts = filtered_df['host'].unique()

                # setup Selenium WebDriver
                driver = setup_driver()
                max_width, max_height = 1920, 1080

                with cProfile.Profile() as pr:
                    # Iterate over each unique host and take a screenshot
                    for index, host in enumerate(unique_hosts[:10]):
                        
                        progress = (index + 1) / len(unique_hosts)
                        progress_bar.progress(progress)
                        status_text.text(f'Numbers of scanned sites: {index + 1}/{len(unique_hosts)}')

                        host, image, error_type = take_screenshot(driver, host, max_width, max_height)
                        if host and image:
                            screenshots.append((host, image))
                        else:
                            errors.append((host, error_type))
                        
                        summary.text(f"Validated Screenshots: {len(screenshots)}, Errors: {len(errors)}")
                with open("profiling_results_screen.txt", "w") as f:
                    stats = pstats.Stats(pr, stream=f)
                    stats.sort_stats('cumulative')
                    stats.print_stats()
                
                # Mostra riepilogo degli errori
                if errors:
                    st.subheader("Errors recap")
                    for host, error_type in errors:
                        st.error(f'Error for \"{host}\": {error_type}')

                driver.quit()
                st.success("Process successfully completed")

                # esistono delle screenshot
                if screenshots:
                    st.subheader("Screenshot gallery")
                    cols = st.columns(3)
                    for i, (host, image) in enumerate(screenshots):
                        with cols[i % 3]:
                            st.image(image, caption=host, use_column_width=True)

                    # Creiamo un file zip in memoria
                    zip_buffer = io.BytesIO()
                    with zipfile.ZipFile(zip_buffer, 'w') as zip_file:
                        for host, image in screenshots:
                            # Convertiamo l'immagine in PNG
                            img_byte_arr = io.BytesIO()
                            image.save(img_byte_arr, format='PNG')
                            img_byte_arr = img_byte_arr.getvalue()
                            
                            # Aggiungiamo l'immagine al file zip
                            zip_file.writestr(f"image_{host}.png", img_byte_arr)

                    # Resettiamo il puntatore del buffer
                    zip_buffer.seek(0)

                    # Utilizziamo st.download_button per creare un pulsante di download
                    st.download_button(
                        label="Download images",
                        data=zip_buffer,
                        file_name="screenshots.zip",
                        mime="application/zip"
                    )


                # Display the screenshots using Streamlit's st.image
                #for host, image in screenshots:
                #    st.image(image, caption=f'Screenshot of {host}', use_column_width=True)

        # Vulnerability Types Analysis
        st.subheader("Top Vulnerability Types")
        vuln_types = filtered_vulnerabilities['template_name'].value_counts().head(10)
        fig_types = px.bar(
            width=_width,
            height=_height,
            x=vuln_types.index, 
            y=vuln_types.values, 
            title="Top 10 Vulnerability Types",
            labels={'x': 'Vulnerability Type', 'y': 'Count'}
        )
        st.plotly_chart(fig_types, use_container_width=True, config={'displayModeBar': False})
        
        types_analysis = ''
        with st.spinner("Analyzing vulnerability types..."):
            types_analysis = ''
            if run_LLM:
                types_analysis = analyze_vulnerability_types(vuln_types.index[0], vuln_types.values[0], vuln_types.index.tolist(), _pipe = pipe, language=language)
        st.markdown(types_analysis)

        # # Remediation Priority Matrix
        # st.header("Remediation Priority Matrix")
        # if all(col in filtered_vulnerabilities.columns for col in [severity_column, 'cvss_score', 'exploit_available']):
        #     fig_remediation = create_severity_impact_bubble(filtered_vulnerabilities, severity_column, 'cvss_score', host_column)
        #     if fig_remediation:
        #         st.plotly_chart(fig_remediation, use_container_width=True, config={'displayModeBar': False})
            
        #     high_priority = filtered_vulnerabilities[(filtered_vulnerabilities['cvss_score'] > 7) & (filtered_vulnerabilities['exploit_available'] == True)]
        #     with st.spinner("Analyzing remediation priorities..."):
        #         remediation_analysis = ''
        #         if run_LLM:
        #             remediation_analysis = analyze_remediation_priority(len(high_priority), total_vulns, _pipe = pipe, language=language)
        #     st.markdown(remediation_analysis)
        # else:
        #     st.info("Not enough information available for remediation priority analysis.")

        st.header('WorldCloud analysis')
        df = load_data_word(file_contents)
        all_tags = df['template_name']
        tag_counts = Counter(all_tags)

        colors = [kelly_green, dodger_blue, burnt_red, mariana_blue]
        n_bins = len(colors)
        cmap_name = 'brand_colors'
        cm = LinearSegmentedColormap.from_list(cmap_name, colors, N=n_bins)

        # Creazione del WordCloud
        wordcloud_ = WordCloud(width=_width, height=_height, 
                            background_color='white', 
                            max_font_size=300, 
                            scale=3, 
                            relative_scaling=0.5, 
                            collocations=False, 
                            colormap=cm).generate_from_frequencies(tag_counts)

        img = wordcloud_.to_image()
        st.image(img, use_column_width=True)

        # Export Options
        st.header("Export Dashboard")
        col1, col2 = st.columns(2)
        with col1:
            export_format = st.selectbox("Choose export format:", ["Word", "CSV", "JSON"], key="export_format")
        with col2:
            if st.button("Generate Report", key="generate_report"):
                with st.spinner(f"Generating {export_format} report..."):
                    analyses = {
                        'overview': overview_analysis,
                        'severity': severity_analysis,
                        'top_vulnerabilities': top_vuln_analysis,
                        'types': types_analysis
                    }
                    if 'cvss_score' in filtered_vulnerabilities.columns:
                        analyses['cvss'] = cvss_analysis
                    #if 'remediation_analysis' in locals():
                    #   analyses['remediation'] = remediation_analysis
                    
                    figures = {
                        'risk_score': fig_risk_score,
                        'severity': fig_severity,
                        'network': fig_network,
                        'types': fig_types
                    }
                    if 'cvss_score' in filtered_vulnerabilities.columns:
                        figures['cvss'] = fig_cvss
                    #if 'fig_remediation' in locals():
                    #   figures['remediation'] = fig_remediation
                    
                    if export_format == "Word":
                        word_buffer = generate_word_report(filtered_vulnerabilities, analyses, figures)
                        st.download_button(
                            label="Download Word Report",
                            data=word_buffer,
                            file_name="orizon_security_report.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        )
                    elif export_format == "CSV":
                        csv = filtered_vulnerabilities.to_csv(index=False)
                        st.download_button(
                            label="Download CSV",
                            data=csv,
                            file_name="vulnerability_report.csv",
                            mime="text/csv",
                        )
                    elif export_format == "JSON":
                        json_str = filtered_vulnerabilities.to_json(orient="records")
                        st.download_button(
                            label="Download JSON",
                            data=json_str,
                            file_name="vulnerability_report.json",
                            mime="application/json",
                        )

        # Interactive Vulnerability Explorer
        st.header("Interactive Vulnerability Explorer")
        
        # Add search functionality
        search_term = st.text_input("Search vulnerabilities", "")
        
        # Select columns to display
        selected_columns = st.multiselect(
            "Select columns to display",
            options=filtered_vulnerabilities.columns,
            default=[host_column, severity_column, 'template_name', 'template_url', description_column],
            key="column_selector"
        )
        
        # Filter vulnerabilities based on search term
        if search_term:
            filtered_data = filtered_vulnerabilities[filtered_vulnerabilities.apply(lambda row: row.astype(str).str.contains(search_term, case=False).any(), axis=1)]
        else:
            filtered_data = filtered_vulnerabilities
        
        # Display filtered data
        #st.dataframe(filtered_data[selected_columns], height=400, use_container_width=True)
        
        # Add pagination
        items_per_page = st.slider("Items per page", min_value=10, max_value=100, value=50, step=10)
        total_pages = len(filtered_data) // items_per_page + (1 if len(filtered_data) % items_per_page > 0 else 0)
        current_page = st.number_input("Page", min_value=1, max_value=total_pages, value=1)
        
        start_idx = (current_page - 1) * items_per_page
        end_idx = start_idx + items_per_page
        st.dataframe(filtered_data[selected_columns].iloc[start_idx:end_idx], height=400, use_container_width=True)
        
        st.write(f"Showing {start_idx+1} to {min(end_idx, len(filtered_data))} of {len(filtered_data)} entries")

    else:
        st.info("Please upload a JSON file in the sidebar to begin the analysis.")

if __name__ == "__main__":
    # login
    name, authentication_status, username = authenticator.login(key='Login', location='main')

    if authentication_status == False:
        st.error('Username/password is incorrect')
    elif authentication_status == None:
        st.warning('Please enter your username and password')
    elif authentication_status:
        # true login
        authenticator.logout('Logout', 'main')
        st.write(f'Welcome *{name}*')
        start = time.time()
        main()
        end = time.time()

        # Calcolo del tempo impiegato
        elapsed_time = end - start

        minutes = int(elapsed_time // 60)
        seconds = int(elapsed_time % 60)

        print(f"Running time: {minutes:.2f}:{seconds:.2f}")